#!/usr/bin/env python3
"""Generate model_benchmarks.docx comparing Claude Fable 5 vs GLM-5.2."""

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Emr as Emt
from docx.enum.table import WdTableAlignment
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# ── Title & header info ───────────────────────────────────────────────
title_el = doc.add_heading(
    "Model Benchmark Comparison: Claude Fable 5 vs GLM-5.2", level=0)

subtitle = doc.add_paragraph("Compiled July 8, 2026")
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

source_note = (
        'Data sourced from BenchLM.ai provisional leaderboards (#1 of 70 models), The AI Rankings,'
        "ComputingForGeeks, and avenchat.com/blog. Arena Elo figures use ± confidence intervals."
)
doc.add_paragraph(source_note).alignment = WD_ALIGN_PARAGRAPH.CENTER

# ── Section: Model Overviews ────────────────────────────────────────
doc.add_heading("1. Model Overview", level=1)


def add_model_para(name, lines):
    p = doc.add_paragraph()
    run = p.runs[0] if p.runs else None

# Claude overview  
p_c_name = doc.add_paragraph()
run_claude_label = p_c_name.add_run("Claude Fable 5")
run_claude_label.bold = True

claude_lines = [
        "Provider: Anthropic",
        "Released: June 9, 2026",
        "Type: Proprietary / Mythos-class MoE (large-scale)",
        "Context Window: 1M+ tokens",
]
for line in claude_lines:
    doc.add_paragraph(line)

doc.add_paragraph(
        "General-purpose model using explicit chain-of-thought reasoning. Production safeguards applied for cybersecurity/biology/chemistry requests."
).italic = True if False else None  # stylistic only — no italic flag on paragraph directly; skip it


# GLM overview  
p_g_name = doc.add_paragraph()
run_glm_label = p_g_name.add_run("GLM-5.2")
run_glm_label.bold = True

glm_lines = [
        "Provider: Zhipu AI (Z.ai)",
        "Released: June 13, 2026",
        'Type: Open-weight MIT license (~753B total parameters)',
        "Context Window: 1M tokens",
]
for line in glm_lines:
    doc.add_paragraph(line)

doc.add_paragraph(
        "Open-weights flagship designed as a coding-focused model with strong open-source availability."
).italic = True if False else None


# ── Section: Summary Comparison Table ────────────────────────  
doc.add_heading("2. Benchmark Comparison — Key Metrics", level=1)

table_data_clean_list = [
        ("Overall Score", "92 / 100 (#1 of 70 provisional)", "~Open-weight category"),

# Arena Elo — text overall 
('Arena Elo — Text', '1509 ± 9.3 (4,350 votes)', '~Mid-to-high tier'),

("Context Window", "1M+ tokens", "1M tokens")


]
table = doc.add_table(rows=len(table_data_clean_list), cols=3)

for i, row in enumerate(table.rows):
    for j in range(3):
        cell = row.cells[j]

# Build the table properly now with clean data from BenchLM.ai + The AI Rankings  
table2_rows_count = len(table_data_clean_list)


doc.add_paragraph('')
